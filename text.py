from urllib.request import urlopen,Request
from urllib.error import HTTPError
from urllib.error import URLError
import urllib.parse
from bs4 import BeautifulSoup
from docx import Document 
import json
import getpass
import re
import sys

def get_input():
    global lang
    global theme
    global soup        

    lang = input('CHOOSE THE LANGUAGE (pt-Português : en-English) --> ')
    
    while True:
        if lang == 'pt' or lang == 'en':

            theme = input('WRITE THE THEME --> ')

            try: 
                url = 'https://'+ lang +'.wikipedia.org/'
                user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.130 Safari/537.36'
                value = {'search': theme}
                headers = {'User-Agent': user_agent}
            except HTTPError as e:
                print(e)
            except URLError:
                print('Server down or incorrect domain')
            else:
                data = urllib.parse.urlencode(value) 
                data = data.encode('utf-8')
                req = Request(url, data, headers)
                resp = urlopen(req) # abrindo url 

                soup = BeautifulSoup(resp, 'html.parser')
            break
        else:
            print(f'Error not found {lang}')
            return get_input()

def decompse(): # Função que remove sujeiras do texto.
    try:
        for links in soup.find_all(['div', 'img', 'h2', 'h3', 'ul', 'sup', 'table', 'span','role'], {'class':['reference', 
            'toc','hatnote', 
            'navbox','mw-indicators mw-body-content',        
            'reflist references-small', 'noprint',
            'reflist columns references-column-width','refbegin reflist columns references-column-width',
            'infobox_v2','vertical-navbox nowraplinks hlist','nowraplinks hlist navbox-inner',
            'infobox biography vcard','infobox',
            'reflist','navigation',
            'box-Não_enciclopédico plainlinks metadata ambox ambox-content ambox-content', 'mw-editsection',
            'mw-indicators mw-body-content','box-Recentism plainlinks metadata ambox ambox-style ambox-Recentism',
            'presentation','official-website','box-Long_plot plainlinks metadata ambox ambox-style ambox-Plot','box-Unreferenced_section plainlinks metadata ambox ambox-content ambox-Unreferenced',
            'metadata plainlinks sistersitebox plainlist mbox-small','navbox authority-control',"mbox-text-span"]}):  # Removendo os links e logos do texto
            links.decompose()

        for refT in soup.find_all(['span','div','h2'], {'id': ['Referências', 'Ligações_externas', 'mw-editsection','mw-headline', 'catlinks','Publica.C3.A7.C3.A3o']}): #Removendo links de referencias
            refT.decompose()

        for st in soup.findAll('table' ,{'style':'border-collapse: collapse'}):
            st.decompose()
    except:pass
    return None

# Funçao que Traz o endereço das imagens. 
def get_link():
 
    href = []

    urls = soup.findAll('img') #pesquisando em todas as tags <img>
    for img in urls:
        if re.findall(r'[jpg]+[png]+', img.get('src')):
            href.append('https:' + img.get('src')) # percorre as tags <img> e tras o contrudo da teg <src>        
        else:
            pass

        link = {'urls': href} # criando um json 

    # salvando em um arquivo .json os links.
    with open('url.json','w') as file:
        json.dump(link,file, indent=2, separators=(',',':'), ensure_ascii=False)
        file.close()
    return None

def get_text():

    usr = getpass.getuser() 
    
    tags = soup.find('div', {'id': 'mw-content-text'})   
    tagsfinal = tags.text.strip()

    if sys.platform.startswith('win32'):
        doc = Document()

        doc.add_heading(theme,0)   
        doc.add_paragraph(tagsfinal)
        doc.save('C:/Users/' + usr + '/Desktop/' + theme + '.docx')
    elif sys.platform.startswith('linux'):
        with open('home/' + usr + '/Área de Trabalho/'+ theme + '.odt','w') as file:
            file.write(tagsfinal)
            file.close()
    else:
        print('System not found !')

    return tagsfinal

def main():   
    get_input()
    decompse()
    get_link()
    get_text()
    
main()